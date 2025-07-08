from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
import os
import subprocess
from docx.shared import Pt, RGBColor
from copy import deepcopy
from docx.text.paragraph import Paragraph
from docx.table import Table



# -----------------------------------------------
def convert_docx_to_pdf(input_path, output_dir=None):
    if not output_dir:
        output_dir = os.path.dirname(input_path)
    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, input_path
        ], check=True)
        print(f"✔ PDF generated in: {output_dir}")
    except Exception as e:
        print("✖ PDF conversion failed:", e)

# -----------------------------------------------
def set_double_bottom_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'BFBFBF')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

def add_header_footer_with_logo(doc_path, output_path, logo_path,
    line1, line2, line3, start_page_number=1, doi_url="", footer_journal="EPRA"):

    doc = Document(doc_path)
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            header._element.remove(para._element)

        table = header.add_table(rows=1, cols=2, width=Inches(8.0))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(0.6)
        table.columns[1].width = Inches(7.4)
        table.autofit = False

        tbl = table._element
        tbl_pr = tbl.xpath(".//w:tblPr")[0]
        borders = OxmlElement('w:tblBorders')
        for b in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'nil')
            borders.append(border)
        tbl_pr.append(borders)

        cell_logo = table.cell(0, 0)
        run = cell_logo.paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(0.4))
        set_double_bottom_border(cell_logo)

        cell_text = table.cell(0, 1)
        para1 = cell_text.paragraphs[0]
        para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run1 = para1.add_run(line1)
        run1.bold = True
        run1.font.size = Pt(9)

        para2 = cell_text.add_paragraph()
        para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = para2.add_run(line2)
        run2.bold = True
        run2.font.size = Pt(11)

        para3 = cell_text.add_paragraph()
        para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run3 = para3.add_run(line3)
        run3.bold = True
        run3.font.size = Pt(8)
        set_double_bottom_border(cell_text)

        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            footer._element.remove(para._element)

        p = footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_text = f"© 2025 {footer_journal} | http://eprajournals.com/"
        if doi_url:
            footer_text += f" | Journal DOI URL: {doi_url}"
        footer_text += " | Page "

        run = p.add_run(footer_text)
        run.font.size = Pt(8)

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    doc.save(output_path)
    print(f"✔ DOCX saved with header/footer at: {output_path}")

# -----------------------------------------------
def is_heading(paragraph):
    text = paragraph.text.strip()
    return text.isupper() and len(text.split()) <= 6 and len(text) > 0

def apply_heading_style(paragraph, style_options, background_color=False):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    run.font.size = Pt(style_options.get("heading_font_size", 11))
    run.font.name = style_options.get("font_name", "Times New Roman")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = paragraph._element
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    if background_color:
        shd.set(qn('w:fill'), style_options.get("heading_bg_color", "E6E6E6"))
    pPr.append(shd)

    rPr = run._element.get_or_add_rPr()
    color = OxmlElement('w:color')
    color.set(qn('w:val'), style_options.get("heading_color", "000000"))
    rPr.append(color)

def apply_font_style(paragraph, style_options):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        run.font.name = style_options.get("font_name", "Times New Roman")
        run.font.size = Pt(style_options.get("font_size", 10))



def is_possible_heading(paragraph):
    text = paragraph.text.strip()
    if not text or not paragraph.runs:
        return None

    is_bold = any(run.bold for run in paragraph.runs) or True
    alignment = paragraph.alignment

    word_count = len(text.split())
    if 1 <= word_count <= 10 and is_bold:
        print(f"Possible heading found: '{text}' with alignment {alignment}")
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return "heading"
        # elif alignment == WD_PARAGRAPH_ALIGNMENT.LEFT or alignment == None:
        else:
            return "subheading"
    return None


# def process_title_author_section(paragraphs):
#     output = []
#     before_abstract = True
#     for para in paragraphs:
#         text = para.text.strip()
#         if not text:
#             continue

#         if text.lower().startswith("abstract"):
#             before_abstract = False
#             break

#         output.append(text)
#     if len(output) >= 3:
#         return {
#             "title": output[0],
#             "authors": output[1],
#             "affiliations": output[2:],
#         }
#     elif len(output) == 2:
#         return {
#             "title": output[0],
#             "authors": output[1],
#             "affiliations": [],
#         }
#     else:
#         return {
#             "title": output[0] if output else "",
#             "authors": "",
#             "affiliations": [],
#         }


def clone_element(el):
    return deepcopy(el)

def is_block_to_remove(text, block):
    return (
        text == block['title']
        or text == block['authors']
        or text == block['corresponding']
        or text in block['affiliations']
    )

def process_title_author_section(paragraphs):
    output = []
    before_abstract = True
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.lower().startswith("abstract"):
            before_abstract = False
            break

        output.append(text)

    block = {
        "title": "",
        "authors": "",
        "affiliations": [],
        "corresponding": ""
    }

    if output:
        block["title"] = output[0]
    if len(output) > 1:
        block["authors"] = output[1]
    if len(output) > 2:
        for line in output[2:]:
            if "corresponding author" in line.lower():
                block["corresponding"] = line
            else:
                block["affiliations"].append(line)

    return block

def style_paragraph(para, font_name, size, bold=False, color=None, align_center=True):
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if align_center else WD_PARAGRAPH_ALIGNMENT.LEFT
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def style_paragraph1(doc, text, font_name, size, bold=False, color=None, align_center=True):
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if align_center else WD_PARAGRAPH_ALIGNMENT.LEFT
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return para

def process_headings(document, heading_font="Times New Roman", heading_size=11, heading_color="000000", subheading_size=10):
    for para in document.paragraphs:
        if "corresponding author" in para.text.lower():
            continue
        tag = is_possible_heading(para)
        # if(tag):``
            # print(f"Processing paragraph:tag = {tag}, text = {para.text.strip()}")
        if tag == "heading":
            # color = "0000FF" if para.text.strip().lower() == "abstract" else None
            style_paragraph(para, heading_font, size = 11, bold=True, color=heading_color, align_center=True)
        elif tag == "subheading":
            style_paragraph(para, heading_font, size = 10, bold=True, color=heading_color, align_center=False)

# ---------- LAYOUTS ----------
def apply_two_column_layout_after_abstract(doc):
    print("Applying two-column layout after ABSTRACT section...")
    abstract_index = -1
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == "ABSTRACT":
            abstract_index = i
            break

    if abstract_index == -1:
        print("No ABSTRACT section found.")
        return

    insert_index = abstract_index + 2
    if insert_index >= len(doc.paragraphs):
        insert_index = len(doc.paragraphs) - 1

    para = doc.paragraphs[insert_index]
    p = para._element
    sectPr = OxmlElement('w:sectPr')
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    sectPr.append(cols)
    p.append(sectPr)

# ---------- MAIN BODY FORMATTING ----------
def process_body_content_with_styles(doc, layout_mode="two_column", style_options={ "font_name": "Times New Roman", "font_size": 10,}):
    found_abstract = False
    found_references = False
    reference_index = 1

    for para in doc.paragraphs:
        text = para.text.strip()

        if text.upper() == "ABSTRACT":
            para.text = "ABSTRACT"
            apply_heading_style(para, style_options, background_color = True)
            found_abstract = True
            continue

        if found_abstract and text != "" and not is_heading(para):
            apply_font_style(para, style_options)
            # found_abstract = False

        elif is_heading(para):
            apply_heading_style(para, style_options, background_color = False)

        elif text.upper() == "REFERENCES":
            para.text = "REFERENCES"
            apply_heading_style(para, style_options, background_color = False)
            found_references = True
            reference_index = 1

        elif found_references and text != "":
            para.text = f"[{reference_index}] {text}"
            reference_index += 1
            apply_font_style(para, style_options)

        else:
            # print(f"Processing paragraph: {text}")
            if found_abstract and not is_heading(para):
                apply_font_style(para, style_options)
            # run.font.size = Pt(10)
            # run.font.name = "Times New Roman"


    if layout_mode == "two_column":
        apply_two_column_layout_after_abstract(doc)
# Apply Two column layout from 1
def apply_two_column_layout(section):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if not cols:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    else:
        cols = cols[0]
    cols.set(qn('w:num'), '2')

def center_tables_and_images(doc):
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Auto-resize columns proportionally if table is too wide
        total_col_width = sum(cell.width or 0 for cell in table.rows[0].cells)
        if total_col_width and total_col_width > usable_width:
            scale = usable_width / total_col_width
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = None  # remove fixed size
    # for para in doc.paragraphs:
    #     for run in para.runs:
    #         if 'graphic' in run._element.xml:
    #             para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #             inline_shapes = para._element.xpath('.//w:drawing//wp:extent')
    #             for extent in inline_shapes:
    #                 cx = int(extent.get('cx'))  # width in EMUs
    #                 max_cx = usable_width.cm * 360000  # cm to EMUs
    #                 if cx > max_cx:
    #                     new_cx = int(max_cx)
    #                     extent.set('cx', str(new_cx))
    for shape in doc.inline_shapes:
        max_width = Inches(6)  # around 6 inches for A4 with 1" margin
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

# -----------------------------------------------
def process_document(input_doc, logo_path, output_doc,
                     journalCode="IJMR",
                     line1="", line2="", line3="",
                     start_page_number=1,
                     doi_url="", footer_journal=""):
    temp_doc = "generated/temp_with_header_7.docx"

    # add_header_footer_with_logo(
    #     doc_path=input_doc,
    #     output_path=temp_doc,
    #     logo_path=logo_path,
    #     line1="ISSN (Online): 2455-3662",
    #     line2="EPRA International Journal of Multidisciplinary Research (IJMR) - Peer Reviewed Journal",
    #     line3="Volume:11 | Issue:6 | June 2025 || Journal DOI: 10.36713/epra2013 || SJIF Impact Factor 2025: 8.691 || ISI Value: 1.188",
    #     start_page_number=3,
    #     doi_url="https://doi.org/10.36713/epra2013",
    #     footer_journal="EPRA IJMR"
    # )
    add_header_footer_with_logo(
        doc_path=input_doc,
        output_path=temp_doc,
        logo_path=logo_path,
        line1=line1,
        line2=line2,
        line3=line3,
        start_page_number=start_page_number,
        doi_url=doi_url,
        footer_journal=footer_journal
    )

    doc = Document(temp_doc)
      # First page title-author block
    # First page title-author block
    block = process_title_author_section(doc.paragraphs)
    print(f"Title: '{block['title']}', Authors: '{block['authors']}', Affiliations: '{block['affiliations']}', Corresponding: '{block['corresponding']}'")
    # Remove old paragraphs matching title/authors/affiliations/corresponding
    new_body = []
    elements_to_keep = []

    for el in doc.element.body:
        # Wrap the XML element as a paragraph or table
        if el.tag.endswith('p'):
            para = Paragraph(el, doc)
            text = para.text.strip()
            if not is_block_to_remove(text, block):
                elements_to_keep.append(clone_element(el))
        elif el.tag.endswith('tbl'):
            elements_to_keep.append(clone_element(el))

    # print("doc.paragraphs:", len(doc.paragraphs))
    seen = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        # if not text:
        #     continue
        # if text in [block['title'], block['authors'], *block['affiliations'], block['corresponding']]:
        #     doc._element.body.remove(para._element)
        if text in [block['title'], block['authors'], block['corresponding']] or text in block['affiliations']:
            continue
        new_body.append(para.text)

    # Clear all paragraphs and rebuild
    doc._body.clear_content()

    # Add styled title, authors, affiliations
    style_paragraph1(doc, block['title'], font_name="Georgia", size=16, bold=True)
    style_paragraph1(doc, block['authors'], font_name="Times New Roman", size=14)
    style_paragraph1(doc, ' '.join(block['affiliations']), font_name="Antiqua", size=11)
    # for aff in block['affiliations']:
    #     style_paragraph1(doc, aff, font_name="Calibri", size=11)
    if block['corresponding']:
        style_paragraph1(doc, block['corresponding'], font_name="Times New Roman", size=10, bold=True)

    # Re-add original body text
    # for text in new_body:
    #     para = doc.add_paragraph(text)
    #     para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for el in elements_to_keep:
        doc.element.body.append(el)

    # for text in new_body:
    #     if text.strip():
    #         para = doc.add_paragraph(text.strip())
    #         para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #         # para.paragraph_format.space_before = Pt(0)
    #         # para.paragraph_format.space_after = Pt(0)

    
    layout_mode = "full_page"  # or "full_page", "two_column"
    style_options = {
        "font_name": "Times New Roman",
        "font_size": 10,
        "heading_font_size": 11,
        "heading_color": "000000",
        "heading_bg_color": "D9D9D9"
    }
    # process_body_content_with_styles(doc, journalCode=journalCode, )
    process_body_content_with_styles(doc, layout_mode=layout_mode, style_options=style_options)
    process_headings(doc)

    # apply_two_column_layout_after_abstract(doc)
    if layout_mode == "two_column":
        for section in doc.sections:
            apply_two_column_layout(section)
        center_tables_and_images(doc)

    doc.save(output_doc)
    print(f"✔ Final document saved at: {output_doc}")
    convert_docx_to_pdf(output_doc)
    

# ------------------- USAGE -------------------
# process_document(
#     input_doc="sample_1.docx",
#     logo_path="logo.png",
#     output_doc="generated/final_output_7.docx",
#     journalCode="IJMR"
# )
